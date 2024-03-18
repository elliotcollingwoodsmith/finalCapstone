"""
!pip install -U pip setuptools wheel
!pip install -U spacy
!pip install spacytextblob
!python -m textblob.download_corpora
!python -m spacy download en_core_web_md
!pip install python-docx
"""

import numpy as np
import pandas as pd
import spacy
import random
import matplotlib.pyplot as plt
import seaborn as sns
from spacytextblob.spacytextblob import SpacyTextBlob
from scipy import stats
from tabulate import tabulate
from urllib.parse import urlparse
from collections import defaultdict
from docx import Document
from docx.shared import Inches

from google.colab import drive
drive.mount('/content/gdrive')
root_path = '/content/gdrive/MyDrive/hyperion/tasks/21_nlp_capstone'

# Load the spaCy model and add the SpacyTextBlob component for sentiment analysis
nlp = spacy.load('en_core_web_md')
nlp.add_pipe('spacytextblob')

# Load the dataset
amazon = pd.read_csv(root_path + '/amazon_product_reviews.csv')

# Display the first 5 rows of the dataset
amazon.head()

# Count the number of non-null values for each column in the DataFrame
non_null_counts = amazon.count()

# Print the counts for each column
print("Number of non-null values for each column:")
print(non_null_counts)

# Select the desired columns for EDA
selected_columns = ['primaryCategories',
                    'name',
                    'sourceURLs',
                    'reviews.rating',
                    'reviews.text',
                    'reviews.title',
                    'reviews.username'
]

# Create a new DataFrame with the selected columns
eda = amazon[selected_columns].copy()

# Drop the "reviews." prefix for consistent variable names
eda.rename(columns=lambda x: x.replace('reviews.', ''), inplace=True)

# Display the new DataFrame
eda.head()

eda.isnull().sum()

# Define a function to extract stem URLs from a list of URLs
def extract_stem_urls(urls):
    # Initialise an empty list to store stem URLs
    stem_urls = []

    # Iterate through each URL in the input list
    for url in urls:
        # Parse the URL to extract its components
        parsed_url = urlparse(url)

        # Split the netloc (domain) of the URL by periods
        netloc_parts = parsed_url.netloc.split('.')

        # Check if the netloc has at least two parts (e.g., 'www.example.com')
        if len(netloc_parts) >= 2:
            # Construct the stem URL using the second-to-last part of the netloc and '.com'
            stem_url = netloc_parts[-2] + '.com'

            # Append the stem URL to the list
            stem_urls.append(stem_url)

    # Return the list of stem URLs
    return stem_urls


# Extract stem URLs from the 'sourceURLs' column
stem_urls = extract_stem_urls(eda['sourceURLs'])

# Create a dictionary to store the counts of reviews per unique stem URL
reviews_per_stem = defaultdict(int)
for stem_url in stem_urls:
    reviews_per_stem[stem_url] += 1

# Find the number of unique stem URLs
num_unique_stem_urls = len(reviews_per_stem)

# Print the number of unique stem URLs
print("Number of Unique Stem URLs:", num_unique_stem_urls)

# Print the number of reviews per stem URL
print("Number of Reviews per Stem URL:")
for stem_url, review_count in reviews_per_stem.items():
    print(f"{stem_url}: {review_count}")

# Compute descriptive statistics for the EDA DataFrame
num_unique_products = eda['name'].nunique()  # Number of unique products
num_unique_users = eda['username'].nunique()  # Number of unique users
num_unique_categories = eda['primaryCategories'].nunique()  # Number of unique categories
total_reviews = eda.shape[0]  # Total number of reviews
avg_reviews_per_user = total_reviews / num_unique_users  # Average reviews per unique user
average_review_length = eda['text'].apply(len).mean()  # Average number of characters per review
average_words_per_review = eda['text'].str.split().apply(len).mean()  # Average number of words per review

# Print the computed statistics
print(f"Number of Unique Products: {num_unique_products}")
print(f"Number of Unique Users: {num_unique_users}")
print(f"Number of Unique Categories: {num_unique_categories}")
print(f"Average Number of Reviews per Unique User: {avg_reviews_per_user:.2f}")
print(f"Average Review Length (Characters): {average_review_length:.2f}")
print(f"Average Words per Review: {average_words_per_review:.2f}")

# List unique products alphabetically
unique_products = sorted(eda['name'].unique())

# List unique categories alphabetically
unique_categories = sorted(eda['primaryCategories'].unique())

# Print unique products
print("Unique Products:")
for product in unique_products:
    print(f"- {product}")

# Print unique categories
print("\nUnique Categories:")
for category in unique_categories:
    print(f"- {category}")

# Ratings variable
ratings = eda['rating']

# Define a function to compute summary statistics for ratings
def numerical_stats(data):
    # Get the total number of reviews
    num_reviews = len(data)

    # Calculate the distribution of ratings
    rating_distribution = ratings.value_counts().sort_index()

    # Calculate the percentage of each rating
    percentage_reviews = (rating_distribution / num_reviews) * 100

    # Create a DataFrame to hold the computed statistics
    stats_df = pd.DataFrame({
        'Stars': rating_distribution.index,
        'Frequency': rating_distribution.values,
        'Percentage': percentage_reviews.values
    })

    # Convert Percentage to string format with one decimal place
    stats_df['Percentage'] = stats_df['Percentage'].apply(lambda x: f"{x:.1f}%")

    return stats_df

# Compute numerical statistics
stats_df = numerical_stats(eda)

# Compute additional statistics using scipy
mean = ratings.mean()
skewness = ratings.skew()
kurtosis = ratings.kurtosis()

# Format rating distribution as a table, excluding the index column
rating_table = tabulate(stats_df, headers='keys', tablefmt='grid', showindex=False)

# Print summary statistics for ratings
print("Summary Statistics for Ratings:")
print(rating_table)
print(f"Mean: {mean:.2f}")
print(f"Skewness: {skewness}")
print(f"Kurtosis: {kurtosis}")

# Histogram of Ratings
plt.figure(figsize=(8, 6))
bin_edges = np.arange(0.5, 6.5, 1)  # Start from 0.5 to align the bars with whole numbers
sns.histplot(ratings, bins=bin_edges, kde=False, color='skyblue', edgecolor='black')
plt.title('Histogram of Ratings')
plt.xlabel('Rating')
plt.ylabel('Frequency')
plt.xticks(range(1, 6))  # Ensure ticks align with whole numbers
plt.grid(axis='y', linestyle='--', alpha=0.7)

# Save the plot as an image
plt.savefig(root_path + '/histogram_of_ratings.png')

plt.show()

# Create a new DataFrame 'cleaned' containing only the 'rating' and 'text' columns from the original DataFrame 'eda'
cleaned = eda[['rating', 'text']].copy()

# Function to preprocess review text data
def preprocess(text):
    doc = nlp(text.lower().strip())
    processed = [token.lemma_.lower() for token in doc if not token.is_stop and not token.is_punct]
    return ' '.join(processed)

# Apply the preprocess function to the 'text' column
cleaned['processed_text'] = cleaned['text'].apply(preprocess)

# Display the first few rows of the cleaned DataFrame
cleaned.head()

# Function to test polarity in cleaned data
def polarity_analysis(processed_text):
    polarities = []
    for index, text in processed_text.items():
        doc = nlp(text)
        polarity_score = doc._.blob.polarity
        polarities.append(polarity_score)
    return polarities

# Add 'polarity' column to the DataFrame
cleaned['polarity'] = polarity_analysis(cleaned['processed_text'])

cleaned.head()

# Function to analyse sentiment based on polarity scores
def sentiment_analysis(polarity_scores):
    sentiment = []
    for polarity_score in polarity_scores:
        if polarity_score > 0:
            sentiment.append('Positive')
        elif polarity_score < 0:
            sentiment.append('Negative')
        else:
            sentiment.append('Neutral')
    return sentiment

# Add 'sentiment' column to the DataFrame using existing polarity scores
cleaned['sentiment'] = sentiment_analysis(cleaned['polarity'])

cleaned.head()

# Compute sentiment statistics
positive_reviews = (cleaned['sentiment'] == 'Positive').sum()
negative_reviews = (cleaned['sentiment'] == 'Negative').sum()
neutral_reviews = (cleaned['sentiment'] == 'Neutral').sum()
average_polarity = cleaned['polarity'].mean()

# Print sentiment statistics
print(f"Number of Positive Reviews: {positive_reviews}")
print(f"Number of Negative Reviews: {negative_reviews}")
print(f"Number of Neutral Reviews: {neutral_reviews}")
print(f"Average Polarity Score: {average_polarity:.2f}")

# Set the figure size for all plots
plt.figure(figsize=(10, 20))  # Increase height for vertical plots

# Bar plot to visualize sentiment distribution
plt.subplot(311)  # Subplot 1
sns.countplot(x='sentiment', hue='sentiment', data=cleaned, palette='pastel')
plt.title('Sentiment Distribution')
plt.xlabel('Sentiment')
plt.ylabel('Frequency')

# Bar plot to visualize sentiment distribution across ratings
plt.subplot(312)  # Subplot 2
sns.countplot(x='rating', hue='sentiment', data=cleaned, palette='pastel', order=[1, 2, 3, 4, 5])
plt.title('Sentiment Distribution Across Ratings')
plt.xlabel('Rating')
plt.ylabel('Frequency')
plt.legend(title='Sentiment')

# Bar plot to visualize sentiment distribution across ratings with a log scale to convey relative proportions
plt.subplot(313)  # Subplot 3
sns.countplot(x='rating', hue='sentiment', data=cleaned, palette='pastel', order=[1, 2, 3, 4, 5])
plt.yscale('log')  # Set log scale for y-axis
plt.title('Sentiment Distribution Across Ratings (Log Scale)')
plt.xlabel('Rating')
plt.ylabel('Frequency (log scale)')
plt.legend(title='Sentiment')

# Adjust layout for better spacing between subplots
plt.subplots_adjust(hspace=0.3)  # Add space between subplots

# Save the combined plot as an image
plt.savefig(root_path + '/sentiment_distribution_plot.png')

plt.show()

# Compute the correlation coefficient between ratings and polarity
correlation = cleaned['rating'].corr(cleaned['polarity'])
print(f"Correlation between ratings and polarity: {correlation}")

# Calculate the average polarity for each rating category
average_polarity_by_rating = cleaned.groupby('rating')['polarity'].mean().reset_index()

# Bar plot to visualize average polarity by ratings
plt.figure(figsize=(10, 6))
sns.barplot(x='rating', y='polarity', data=average_polarity_by_rating, color='skyblue')
plt.title('Average Polarity by Ratings')
plt.xlabel('Rating')
plt.ylabel('Average Polarity')

# Add vertical line at y=0 using Matplotlib
plt.axhline(y=0, color='red', linestyle='--', linewidth=1)

# Save the plot as an image
plt.savefig(root_path + '/average_polarity_plot.png')

plt.show()

# Box plot to compare the distribution of polarity scores for each rating category
plt.figure(figsize=(10, 6))
sns.boxplot(x='rating', y='polarity', data=cleaned, palette='pastel', hue='rating', legend=False)
plt.title('Distribution of Polarity Scores Across Ratings')
plt.xlabel('Rating')
plt.ylabel('Polarity Score')

# Save the plot as an image
plt.savefig(root_path + '/boxplot_polarity_scores.png')

plt.show()

# Function to test similarity between two reviews
def similarity(first, second):
    similarity_result = nlp(first).similarity(nlp(second))
    return similarity_result

# Select two random indices
first_index = random.randint(0, len(cleaned) - 1)
second_index = random.randint(0, len(cleaned) - 1)

# Get reviews based on random indices
first_review = cleaned['text'].iloc[first_index]
second_review = cleaned['text'].iloc[second_index]

# Print similarity between two randomly chosen reviews
print(f"\nReview One: {first_review}")
print(f"Review Two: {second_review}")
print(f"Similarity: {similarity(first_review, second_review)}")

# Write detailed report
doc = Document()

doc.add_heading('Sentiment Analysis Report', level=1)

doc.add_heading('1. Dataset Description and Additional Descriptive Statistics:', level=2)
doc.add_paragraph(f"- The dataset used in this analysis comprises {len(eda)} Amazon product reviews.")
doc.add_paragraph(f"- The reviews cover {num_unique_products} unique products sourced from {num_unique_stem_urls} unique websites, providing a comprehensive sample of consumer feedback.")

doc.add_heading('2. Descriptive Statistics and Distribution of Ratings:', level=2)
doc.add_paragraph(f"- Number of Unique Users: {num_unique_users}")
doc.add_paragraph(f"- Average Number of Reviews per Unique User: {avg_reviews_per_user:.2f}")
doc.add_paragraph(f"- Average Review Length: {average_words_per_review:.2f} words")
doc.add_paragraph(f"- Average Rating: {mean:.2f}")
doc.add_paragraph('   - Distribution Insights:')
doc.add_paragraph('       - Skewness:')
doc.add_paragraph('           - Measures the symmetry of the distribution of ratings.')
doc.add_paragraph('           - A skewness value close to 0 indicates a symmetrical distribution')
doc.add_paragraph('           - Positive skewness indicates a right-skewed distribution.')
doc.add_paragraph('           - Negative skewness indicates a left-skewed distribution.')
doc.add_paragraph(f'       - The distribution of ratings in our dataset exhibits {skewness:.2f} skewness and {kurtosis:.2f} kurtosis. The positive skewness suggests a right-skewed distribution, indicating that a majority of the reviews are positive.')

doc.add_heading('3. Justification for Sentiment Analysis:', level=2)
doc.add_paragraph('   - The skewness and kurtosis values provide insights into the distribution of ratings in the dataset. By contrast, sentiment analysis offers a qualitative assessment of the sentiment expressed in reviews.')
doc.add_paragraph('   - While numerical statistics such as skewness and kurtosis provide a quantitative overview, sentiment analysis delves deeper into the emotional content of the reviews, offering a nuanced understanding of consumer sentiment.')

doc.add_heading('4. Preprocessing Steps:', level=2)
doc.add_paragraph('   - Text data preprocessing is a crucial step in preparing the reviews for sentiment analysis.')
doc.add_paragraph('   - The following preprocessing steps were applied:')
doc.add_paragraph('       - Converting text to lowercase to ensure consistency in analysis.')
doc.add_paragraph('       - Stripping whitespace to remove unnecessary characters.')
doc.add_paragraph('       - Lemmatization to reduce words to their base form for better analysis.')
doc.add_paragraph('       - Removing stopwords and punctuation to focus on meaningful content.')
doc.add_paragraph('   - These steps help improve the quality of the input data and enhance the accuracy of sentiment analysis.')

doc.add_heading('5. Sentiment Analysis Implementation Process:', level=2)
doc.add_paragraph('   - Model Selection: The sentiment analysis model was implemented using the spaCy library with the textblob extension, chosen for its capabilities in natural language processing tasks and sentiment analysis.')
doc.add_paragraph('   - Model Application: The sentiment analysis model was applied to each review in the dataset using the `analyze_sentiment` function, which determined the sentiment based on the polarity score provided by the textblob extension.')
doc.add_paragraph('   - Evaluation of Results: The performance of the sentiment analysis model was evaluated by computing sentiment statistics, including the number of positive, negative, and neutral reviews, as well as the average sentiment polarity across all reviews.')
doc.add_paragraph('   - Visualization of Results: Bar plots were created to visualize the distribution of sentiments in the dataset, providing a clear overview of the sentiment expressed in the product reviews and the relationship between sentiment and ratings.')

doc.add_heading('6. Evaluation of Results:', level=2)
doc.add_paragraph('   - Quantitative Evaluation:')
doc.add_paragraph(f'       - Number of reviews analyzed: {len(eda)}')
doc.add_paragraph('       - Distribution of sentiment:')
doc.add_paragraph(f'           - Positive: {positive_reviews}')
doc.add_paragraph(f'           - Negative: {negative_reviews}')
doc.add_paragraph(f'           - Neutral: {neutral_reviews}')
doc.add_paragraph(f'       - Average sentiment polarity: {average_polarity:.2f}')
doc.add_paragraph(f'       - Correlation between ratings and polarity: {correlation:.2f}')
doc.add_paragraph('   - Analysis:')
doc.add_paragraph(f'       - The dataset contains a total of {len(eda)} product reviews from Amazon. The average review length is {average_words_per_review:.2f} words.')
doc.add_paragraph(f'       - The sentiment analysis reveals that there are {positive_reviews} positive reviews, {negative_reviews} negative reviews, and {neutral_reviews} neutral reviews.')
doc.add_paragraph(f'       - Polarity scores range from -1 to 1, with positive values indicating positive sentiment, negative values indicating negative sentiment, and values closer to 0 indicating neutrality.')
doc.add_paragraph(f'       - The average sentiment polarity across all reviews is {average_polarity:.2f}, indicating an overall positive sentiment.')
doc.add_paragraph(f'       - The correlation coefficient between ratings and polarity is 0.33, suggesting a weak to moderately positive correlation.')
doc.add_paragraph(f'       - This may indicate that a reviewer\'s sentiment does not strongly correlate with the star rating they assign, which could be surprising given that both metrics are reflective of the same reviewer\'s opinion.')

doc.add_heading('7. Visualization of Review Ratings and Sentiment Analysis:', level=2)
doc.add_paragraph('   - Histogram of Ratings:')
doc.add_picture(root_path + '/histogram_of_ratings.png', width=Inches(5))
doc.add_paragraph('   - Bar plot of Sentiment Distribution:')
doc.add_picture(root_path + '/sentiment_distribution_plot.png', width=Inches(5))

doc.add_heading('8. Distribution of Polarity Scores:', level=2)
doc.add_paragraph('   - Box plot of Polarity Scores:')
doc.add_picture(root_path + '/boxplot_polarity_scores.png', width=Inches(5))

doc.add_heading('9. Distribution of Polarity Scores Across Ratings:', level=2)
doc.add_picture(root_path + '/average_polarity_plot.png', width=Inches(5))

doc.add_heading('10. Review Similarity Analysis:')
doc.add_paragraph('   - Calculated similarity scores between pairs of reviews using spaCy\'s similarity comparison functionality.')
doc.add_paragraph('   - Scores range from 0 to 1, with higher values indicating greater similarity.')
doc.add_paragraph('   - High scores suggest shared content or sentiment, while lower scores indicate differences.')
doc.add_paragraph('   - Example pairs were randomly selected to illustrate the analysis:')
doc.add_paragraph('       - Review One: "It\'s not a high discharge battery, but for typical uses such as wireless mouse and flashlights, these are great."')
doc.add_paragraph('       - Review Two: "The Kindle is easiest to use, with graphics and screen crisp, clear, and brilliant colors."')
doc.add_paragraph('   - The similarity score between these reviews is 0.89, indicating a high degree of similarity in content or sentiment.')
doc.add_paragraph('   - Note that scores reflect textual similarity and may not capture semantic context comprehensively.')

doc.add_heading('11. Insights into Model\'s Strengths and Limitations:')
doc.add_paragraph('   - Strengths:')
doc.add_paragraph('       - The model effectively predicts sentiment based on polarity scores, providing valuable insights into consumer sentiment.')
doc.add_paragraph('       - Preprocessing steps enhance the quality of input data, leading to more accurate sentiment analysis results.')
doc.add_paragraph('   - Limitations:')
doc.add_paragraph('       - The model\'s accuracy may vary depending on the complexity of language and use of slang in reviews.')
doc.add_paragraph('       - In some cases, the sentiment scores may not accurately reflect the sentiment of the review, indicating areas for improvement.')
doc.add_paragraph('       - Comparing results with other sentiment analysis methods like VADER could provide additional insights and enhance reliability.')
doc.add_paragraph('       - The imbalanced distribution of positive reviews may bias the model towards positive sentiments, potentially affecting the accuracy of sentiment classification.')

doc.add_heading('12. Recommendations for Improvement:')
doc.add_paragraph('   - Comparing sentiment analysis results with alternative methods like VADER to assess reliability and improve accuracy.')
doc.add_paragraph('   - Explore techniques for handling slang and poorly written sentences to enhance the model\'s effectiveness in capturing nuanced sentiment expressions.')
doc.add_paragraph('   - Continuously evaluate and refine the sentiment analysis model based on feedback and evolving language patterns.')
doc.add_paragraph('   - Optimise the model\'s hyperparameters to maximise performance.')
doc.add_paragraph('   - Evaluate the trained model\'s performance on a separate test set using various evaluation metrics.')
doc.add_paragraph('   - Deploy the model into production and monitor its performance over time for continuous improvement.')

# Path to save the Word document
doc_path = root_path + '/sentiment_analysis_report.docx'

# Save report to the root path
doc.save(doc_path)

print(f"Word document report saved at: {doc_path}")